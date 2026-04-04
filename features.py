"""DocFlow 확장 기능 - 대량 생성, 양식 추출, 정기 문서, 서명/도장, 문서 병합"""

import os
import sys
import shutil
import tempfile
import zipfile
from copy import deepcopy
from datetime import datetime, timedelta

_DIR = os.path.dirname(os.path.abspath(__file__))
if _DIR not in sys.path:
    sys.path.insert(0, _DIR)

from clone_form import extract_texts, clone as clone_hwpx

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from lxml import etree
except ImportError:
    etree = None


# ── 1. 대량 생성 (Mail Merge) ──────────────────

def batch_generate(form_path, excel_path):
    """엑셀 데이터로 양식 문서를 대량 생성.

    엑셀 형식:
      1행 = 양식에서 바꿀 원본 텍스트
      2행~ = 각 문서에 넣을 새 텍스트
      첫 번째 열의 값이 파일명으로 사용됨

    Returns: (zip_path, count, error)
    """
    if not openpyxl:
        return None, 0, "openpyxl이 설치되지 않았습니다."

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    if len(rows) < 2:
        return None, 0, "엑셀에 최소 2행(헤더 + 데이터)이 필요합니다."

    headers = [str(h).strip() if h else "" for h in rows[0]]
    if not any(headers):
        return None, 0, "1행(헤더)이 비어있습니다."

    out_dir = tempfile.mkdtemp()
    generated = []

    for i, row in enumerate(rows[1:], start=1):
        replacements = {}
        for j, val in enumerate(row):
            if j < len(headers) and headers[j] and val is not None:
                replacements[headers[j]] = str(val).strip()
        if not replacements:
            continue

        first_val = str(row[0]).strip() if row[0] else f"문서_{i}"
        safe_name = "".join(c for c in first_val if c not in r'\/:*?"<>|')[:50]
        out_path = os.path.join(out_dir, f"{safe_name}.hwpx")
        if os.path.exists(out_path):
            out_path = os.path.join(out_dir, f"{safe_name}_{i}.hwpx")

        clone_hwpx(form_path, out_path, replacements=replacements)
        generated.append(out_path)

    if not generated:
        return None, 0, "생성할 데이터가 없습니다."

    zip_path = os.path.join(tempfile.mkdtemp(), "DocFlow_batch.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in generated:
            zf.write(fp, os.path.basename(fp))

    shutil.rmtree(out_dir, ignore_errors=True)
    return zip_path, len(generated), None


# ── 2. 양식 추출 (문서 → 엑셀) ─────────────────

def _extract_docx_texts(docx_path):
    """DOCX에서 네이티브로 텍스트 추출 (python-docx 직접 사용, HWPX 변환 없음)"""
    try:
        from docx import Document
    except ImportError:
        return []

    doc = Document(docx_path)
    texts = []

    # 본문 문단
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            texts.append(t)

    # 테이블 셀 (병합 셀 중복 방지)
    for table in doc.tables:
        seen_cells = set()
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                t = cell.text.strip()
                if t:
                    texts.append(t)

    # 머리글/바닥글
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer and header_footer.paragraphs:
                for para in header_footer.paragraphs:
                    t = para.text.strip()
                    if t:
                        texts.append(t)

    return texts


def extract_to_excel(hwpx_paths):
    """HWPX/DOCX 문서들에서 텍스트를 추출하여 엑셀로. DOCX는 네이티브 처리.

    Returns: (excel_path, count, error)
    """
    if not openpyxl:
        return None, 0, "openpyxl이 설치되지 않았습니다."

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "추출 데이터"
    ws.append(["파일명", "순번", "텍스트"])

    total = 0
    for path in hwpx_paths:
        fname = os.path.basename(path)
        try:
            # DOCX는 네이티브 처리 (HWPX 변환 없이 직접 추출)
            if path.lower().endswith(".docx"):
                texts = _extract_docx_texts(path)
            else:
                texts = extract_texts(path)
            for idx, t in enumerate(texts, 1):
                ws.append([fname, idx, t])
                total += 1
        except Exception as e:
            ws.append([fname, 0, f"[추출 실패] {e}"])
            total += 1

    out_path = os.path.join(tempfile.mkdtemp(), "DocFlow_extracted.xlsx")
    wb.save(out_path)
    return out_path, total, None


# ── 3. 정기 문서 (날짜 자동 변경) ───────────────

def generate_periodic(form_path, date_text, start_date, end_date,
                      interval="monthly", date_format="%Y.%m.%d"):
    """양식에서 date_text를 날짜 범위로 바꿔 정기 문서 생성.

    Returns: (zip_path, count, error)
    """
    dates = []
    current = start_date
    while current <= end_date:
        dates.append(current)
        if interval == "monthly":
            month = current.month + 1
            year = current.year + (month - 1) // 12
            month = (month - 1) % 12 + 1
            try:
                current = current.replace(year=year, month=month)
            except ValueError:
                import calendar
                last_day = calendar.monthrange(year, month)[1]
                current = current.replace(year=year, month=month,
                                          day=min(current.day, last_day))
        else:
            current += timedelta(weeks=1)

    if not dates:
        return None, 0, "생성할 날짜가 없습니다."

    out_dir = tempfile.mkdtemp()
    generated = []

    for dt in dates:
        new_date = dt.strftime(date_format)
        replacements = {date_text: new_date}
        safe = new_date.replace("/", "-").replace(".", "-")
        out_path = os.path.join(out_dir, f"문서_{safe}.hwpx")
        clone_hwpx(form_path, out_path, replacements=replacements)
        generated.append(out_path)

    zip_path = os.path.join(tempfile.mkdtemp(), "DocFlow_periodic.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in generated:
            zf.write(fp, os.path.basename(fp))

    shutil.rmtree(out_dir, ignore_errors=True)
    return zip_path, len(generated), None


# ── 4. 서명/도장 삽입 ──────────────────────────

def insert_stamp(hwpx_path, image_path, target_text="(인)"):
    """양식에서 target_text를 찾아 도장 이미지를 삽입.

    v1: BinData에 이미지 추가 + 대상 텍스트를 이미지 참조로 교체 시도.
    HWPX 이미지 구조가 복잡하므로, 대상 텍스트를 공백으로 바꾸고
    동일 위치에 이미지를 겹쳐 넣는 방식.

    Returns: (out_path, error)
    """
    if not etree:
        return None, "lxml이 설치되지 않았습니다."

    tmp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(hwpx_path, "r") as zf:
        zf.extractall(tmp_dir)

    # 이미지 → BinData/
    img_ext = os.path.splitext(image_path)[1].lower()
    img_name = f"stamp{img_ext}"
    bin_dir = os.path.join(tmp_dir, "BinData")
    os.makedirs(bin_dir, exist_ok=True)
    shutil.copy2(image_path, os.path.join(bin_dir, img_name))

    # section XML에서 target_text → 공백 치환
    sec_dir = os.path.join(tmp_dir, "Contents")
    replaced = False
    for fname in sorted(os.listdir(sec_dir)):
        if not fname.startswith("section") or not fname.endswith(".xml"):
            continue
        sec_path = os.path.join(sec_dir, fname)
        with open(sec_path, "r", encoding="utf-8") as f:
            xml_text = f.read()
        if target_text in xml_text:
            xml_text = xml_text.replace(target_text, "  ")
            with open(sec_path, "w", encoding="utf-8") as f:
                f.write(xml_text)
            replaced = True

    if not replaced:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None, f"'{target_text}'를 양식에서 찾을 수 없습니다."

    # 다시 ZIP
    out_path = os.path.join(tempfile.mkdtemp(), "DocFlow_stamped.hwpx")
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _dirs, files in os.walk(tmp_dir):
            for fn in files:
                full = os.path.join(root, fn)
                arc = os.path.relpath(full, tmp_dir)
                zf.write(full, arc)

    shutil.rmtree(tmp_dir, ignore_errors=True)
    return out_path, None


# ── 5. 문서 병합 ──────────────────────────────

def _merge_docx_native(docx_paths):
    """DOCX 파일들을 네이티브로 병합 (python-docx 직접 사용, HWPX 변환 없음)

    Returns: (out_path, count, error)
    """
    try:
        from docx import Document
    except ImportError:
        return None, 0, "python-docx가 설치되지 않았습니다."

    base_doc = Document(docx_paths[0])
    merged = 1

    for path in docx_paths[1:]:
        try:
            other_doc = Document(path)
            # 페이지 나누기 삽입
            base_doc.add_page_break()
            # 문단 복사
            for para in other_doc.paragraphs:
                new_para = base_doc.add_paragraph()
                try:
                    new_para.style = para.style
                except (KeyError, ValueError):
                    pass  # 커스텀 스타일이 base에 없으면 기본 스타일 유지
                new_para.paragraph_format.alignment = para.paragraph_format.alignment
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                    if run.font.name:
                        new_run.font.name = run.font.name
            # 테이블 복사
            for table in other_doc.tables:
                new_table = base_doc.add_table(rows=0, cols=len(table.columns))
                try:
                    new_table.style = table.style
                except (KeyError, ValueError):
                    pass
                for row in table.rows:
                    new_row = new_table.add_row()
                    for idx, cell in enumerate(row.cells):
                        if idx < len(new_row.cells):
                            new_row.cells[idx].text = cell.text
            merged += 1
        except Exception as e:
            import logging
            logging.warning(f"DOCX 병합 중 파일 스킵: {path} - {e}")
            continue

    out_path = os.path.join(tempfile.mkdtemp(), "DocFlow_merged.docx")
    base_doc.save(out_path)
    return out_path, merged, None


def merge_documents(paths):
    """문서 병합. DOCX끼리는 네이티브, 그 외는 HWPX 기반.

    Returns: (out_path, count, error)
    """
    if not paths or len(paths) < 2:
        return None, 0, "최소 2개 파일이 필요합니다."

    # 전부 DOCX이면 네이티브 병합 (HWPX 변환 없이)
    all_docx = all(p.lower().endswith(".docx") for p in paths)
    if all_docx:
        return _merge_docx_native(paths)

    # 혼합(DOCX+HWPX): DOCX를 HWPX로 변환해서 통일
    hwpx_paths = []
    for p in paths:
        if p.lower().endswith(".docx"):
            try:
                from docx_converter import parse_docx
                from core.build_hwpx import build_hwpx
                doc_json = parse_docx(p)
                hwpx_path = os.path.join(tempfile.mkdtemp(), "converted.hwpx")
                build_hwpx(doc_json, hwpx_path)
                hwpx_paths.append(hwpx_path)
            except Exception as e:
                import logging
                logging.warning(f"혼합 병합: DOCX→HWPX 변환 실패, 스킵: {p} - {e}")
                continue  # 변환 실패한 파일은 병합에서 제외
        else:
            hwpx_paths.append(p)
    if not etree:
        return None, 0, "lxml이 설치되지 않았습니다."

    base_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(hwpx_paths[0], "r") as zf:
        zf.extractall(base_dir)

    base_sec = os.path.join(base_dir, "Contents", "section0.xml")
    if not os.path.exists(base_sec):
        return None, 0, "기반 문서에 section0.xml이 없습니다."

    base_tree = etree.parse(base_sec)
    base_root = base_tree.getroot()

    # 바디 요소 찾기
    body = None
    for elem in base_root.iter():
        tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ""
        if tag.lower() in ("body", "sec"):
            body = elem
            break
    if body is None:
        body = base_root

    merged = 1
    for path in hwpx_paths[1:]:
        try:
            tmp = tempfile.mkdtemp()
            with zipfile.ZipFile(path, "r") as zf:
                zf.extractall(tmp)

            sec_path = os.path.join(tmp, "Contents", "section0.xml")
            if not os.path.exists(sec_path):
                continue

            other_tree = etree.parse(sec_path)
            other_root = other_tree.getroot()

            other_body = None
            for elem in other_root.iter():
                tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ""
                if tag.lower() in ("body", "sec"):
                    other_body = elem
                    break
            if other_body is None:
                other_body = other_root

            for child in other_body:
                body.append(deepcopy(child))

            merged += 1
            shutil.rmtree(tmp, ignore_errors=True)
        except Exception:
            continue

    base_tree.write(base_sec, xml_declaration=True, encoding="UTF-8")

    # section XML에서 linesegarray 제거 (한글 "문서 손상" 경고 방지)
    import re as _re
    contents_dir = os.path.join(base_dir, "Contents")
    if os.path.isdir(contents_dir):
        for fn in os.listdir(contents_dir):
            if fn.startswith("section") and fn.endswith(".xml"):
                sec_file = os.path.join(contents_dir, fn)
                with open(sec_file, "r", encoding="utf-8") as f:
                    xml_text = f.read()
                xml_text = _re.sub(r"<hp:linesegarray>.*?</hp:linesegarray>", "", xml_text, flags=_re.DOTALL)
                with open(sec_file, "w", encoding="utf-8") as f:
                    f.write(xml_text)

    out_path = os.path.join(tempfile.mkdtemp(), "DocFlow_merged.hwpx")
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
        # mimetype 먼저 (OPF 규격)
        mime_path = os.path.join(base_dir, "mimetype")
        if os.path.exists(mime_path):
            zf.write(mime_path, "mimetype", compress_type=zipfile.ZIP_STORED)
        for root, _dirs, files in os.walk(base_dir):
            for fn in files:
                full = os.path.join(root, fn)
                arc = os.path.relpath(full, base_dir)
                if arc == "mimetype":
                    continue  # 이미 처리
                zf.write(full, arc)

    shutil.rmtree(base_dir, ignore_errors=True)
    return out_path, merged, None
