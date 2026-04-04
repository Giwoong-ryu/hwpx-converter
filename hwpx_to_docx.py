"""HWPX → DOCX 변환 모듈
HWPX(ZIP+XML)에서 텍스트와 테이블을 추출하여 python-docx로 DOCX를 생성한다.
"""

import re
import zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_hwpx_to_docx(hwpx_path: str, docx_path: str) -> str:
    """HWPX 파일을 DOCX로 변환한다.

    Args:
        hwpx_path: 입력 HWPX 파일 경로
        docx_path: 출력 DOCX 파일 경로

    Returns:
        출력 파일 경로
    """
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    with zipfile.ZipFile(hwpx_path, "r") as zf:
        for name in sorted(zf.namelist()):
            if not (name.startswith("Contents/section") and name.endswith(".xml")):
                continue

            data = zf.read(name).decode("utf-8")

            # <hp:p> 문단 단위로 처리
            paragraphs = re.findall(r"<hp:p\b[^>]*>(.*?)</hp:p>", data, re.DOTALL)

            for para_xml in paragraphs:
                # 테이블 안의 텍스트는 별도 처리하므로 스킵
                if "<hp:tbl" in para_xml:
                    continue

                # <hp:t> 태그에서 텍스트 추출
                texts = re.findall(r"<hp:t>(.*?)</hp:t>", para_xml, re.DOTALL)
                if not texts:
                    doc.add_paragraph("")
                    continue

                # 인라인 XML 태그 제거
                clean_texts = []
                for t in texts:
                    clean = re.sub(r"<[^>]+>", "", t).strip()
                    if clean:
                        clean_texts.append(clean)

                if not clean_texts:
                    doc.add_paragraph("")
                    continue

                line = " ".join(clean_texts)
                p = doc.add_paragraph(line)

                # 정렬 감지
                if 'jc="CENTER"' in para_xml or 'jc="center"' in para_xml:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'jc="RIGHT"' in para_xml or 'jc="right"' in para_xml:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 테이블 처리
            tables = re.findall(r"<hp:tbl\b[^>]*>(.*?)</hp:tbl>", data, re.DOTALL)
            for tbl_xml in tables:
                rows_xml = re.findall(r"<hp:tr\b[^>]*>(.*?)</hp:tr>", tbl_xml, re.DOTALL)
                if not rows_xml:
                    continue

                # 첫 행에서 열 수 파악
                first_cells = re.findall(r"<hp:tc\b[^>]*>(.*?)</hp:tc>", rows_xml[0], re.DOTALL)
                n_cols = max(len(first_cells), 1)

                table = doc.add_table(rows=0, cols=n_cols)
                table.style = "Table Grid"

                for row_xml in rows_xml:
                    cells_xml = re.findall(r"<hp:tc\b[^>]*>(.*?)</hp:tc>", row_xml, re.DOTALL)
                    row = table.add_row()

                    for ci, cell_xml in enumerate(cells_xml):
                        if ci >= n_cols:
                            break
                        cell_texts = re.findall(r"<hp:t>(.*?)</hp:t>", cell_xml, re.DOTALL)
                        clean = " ".join(re.sub(r"<[^>]+>", "", t).strip() for t in cell_texts)
                        row.cells[ci].text = clean

    doc.save(docx_path)
    return docx_path
