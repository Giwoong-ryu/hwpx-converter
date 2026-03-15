"""
HWPX 양식 채우기 모듈
기존 .hwpx 파일의 {{플레이스홀더}}를 데이터로 치환
"""
import os
import re
import uuid
import tempfile
from pathlib import Path
from zipfile import ZipFile, ZIP_STORED, ZIP_DEFLATED

from lxml import etree
from page_guard import collect_metrics, compare_metrics

PLACEHOLDER_RE = re.compile(r'\{\{([^}]+)\}\}')


def analyze_template(hwpx_path: str) -> tuple[list[str], str]:
    """HWPX 파일에서 {{placeholder}} 목록 추출.

    Returns:
        (필드명 리스트, 요약 텍스트)
    """
    try:
        with ZipFile(hwpx_path, 'r') as zf:
            if 'Contents/section0.xml' not in zf.namelist():
                return [], "section0.xml이 없는 파일입니다."
            section_xml = zf.read('Contents/section0.xml').decode('utf-8')
    except Exception as e:
        return [], f"HWPX 파일 읽기 실패: {e}"

    placeholders = []
    seen = set()
    for match in PLACEHOLDER_RE.finditer(section_xml):
        name = match.group(1).strip()
        if name not in seen:
            placeholders.append(name)
            seen.add(name)

    if not placeholders:
        summary = "{{플레이스홀더}}를 찾지 못했습니다.\n"
        summary += "양식 파일에 {{이름}}, {{부서}} 등의 마커를 넣어주세요."
    else:
        summary = f"검출된 필드 {len(placeholders)}개:\n"
        summary += ", ".join(f"{{{{{p}}}}}" for p in placeholders)
        summary += "\n\n아래 JSON에 값을 입력하세요:\n"
        example = {p: "" for p in placeholders}
        import json
        summary += json.dumps(example, ensure_ascii=False, indent=2)

    return placeholders, summary


def fill_template(hwpx_path: str, values: dict) -> tuple[str, str]:
    """HWPX 파일의 {{placeholder}}를 values로 치환한 새 파일 생성.

    Returns:
        (출력 파일 경로, 로그 텍스트)
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        work_dir = os.path.join(tmpdir, 'work')

        # 1. Unzip
        with ZipFile(hwpx_path, 'r') as zf:
            zf.extractall(work_dir)

        # 2. section0.xml 치환
        section_path = os.path.join(work_dir, 'Contents', 'section0.xml')
        with open(section_path, 'r', encoding='utf-8') as f:
            content = f.read()

        replaced_count = 0
        for key, value in values.items():
            safe_value = _xml_escape(str(value))
            marker = f'{{{{{key}}}}}'
            count = content.count(marker)
            if count > 0:
                content = content.replace(marker, safe_value)
                replaced_count += count

        # 치환 안 된 플레이스홀더 확인
        remaining = PLACEHOLDER_RE.findall(content)

        with open(section_path, 'w', encoding='utf-8') as f:
            f.write(content)

        # 3. XML 검증
        try:
            etree.fromstring(content.encode('utf-8'))
        except etree.XMLSyntaxError as e:
            raise RuntimeError(f"치환 후 XML 검증 실패: {e}")

        # 4. Repack
        unique = uuid.uuid4().hex[:8]
        output_name = f"filled_{unique}.hwpx"
        output_path = os.path.join(tempfile.gettempdir(), output_name)

        _pack_hwpx(Path(work_dir), Path(output_path))

        # 5. page_guard 검증 (원본 대비 구조 비교)
        from pathlib import Path as _Path
        ref_metrics = collect_metrics(_Path(hwpx_path))
        out_metrics = collect_metrics(_Path(output_path))
        guard_errors = compare_metrics(ref_metrics, out_metrics, 0.15, 0.25)

        # 6. 로그
        log = f"치환 완료\n"
        log += f"  치환된 필드: {replaced_count}개\n"
        if remaining:
            log += f"  미치환 필드: {', '.join(f'{{{{{r}}}}}' for r in remaining)}\n"
        log += f"  파일: {output_name}\n"
        if guard_errors:
            log += f"\n[page_guard 경고] {len(guard_errors)}건\n"
            for ge in guard_errors:
                log += f"  - {ge}\n"
        else:
            log += "\n[page_guard] 통과 - 원본 구조와 일치"

        return output_path, log


def parse_values_file(file_path: str) -> dict:
    """Excel/CSV/JSON 파일에서 key-value 쌍 추출.

    지원 형식:
    - Excel/CSV: A열=필드명, B열=값 (2열 구조)
    - JSON: {"필드명": "값", ...} 객체
    """
    import json as _json
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.json':
        with open(file_path, 'r', encoding='utf-8') as f:
            data = _json.load(f)
        if isinstance(data, dict):
            return {str(k): str(v) for k, v in data.items()}
        raise ValueError("JSON 파일은 {필드: 값} 객체 형식이어야 합니다")

    if ext == '.csv':
        import csv
        values = {}
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) >= 2 and row[0].strip():
                    values[row[0].strip()] = row[1].strip()
        if not values:
            raise ValueError("CSV 파일에서 값을 찾지 못했습니다. A열=필드명, B열=값 형식으로 작성하세요.")
        return values

    if ext in ('.xlsx', '.xls'):
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        ws = wb.active
        values = {}
        for row in ws.iter_rows(min_col=1, max_col=2):
            if len(row) >= 2:
                key = str(row[0].value or '').strip()
                val = str(row[1].value or '').strip()
                if key and key != 'None':
                    values[key] = val if val != 'None' else ''
        wb.close()
        if not values:
            raise ValueError("Excel 파일에서 값을 찾지 못했습니다. A열=필드명, B열=값 형식으로 작성하세요.")
        return values

    if ext in ('.docx',):
        return _extract_docx_text_as_values(file_path)

    if ext == '.hwpx':
        return _extract_hwpx_text_as_values(file_path)

    if ext == '.txt':
        return _parse_txt_keyvalue(file_path)

    raise ValueError(f"지원하지 않는 파일 형식: {ext} (Excel, CSV, JSON, DOCX, HWPX, TXT 지원)")


def _extract_docx_text_as_values(docx_path: str) -> dict:
    """DOCX 파일에서 텍스트 추출 → 필드명:값 매핑.

    테이블 구조: 왼쪽 셀=필드명, 오른쪽 셀=값.
    테이블 없으면: 각 문단을 "문단1", "문단2" 키로 매핑.
    """
    from docx import Document

    doc = Document(docx_path)
    values = {}

    # 테이블에서 key-value 추출 시도
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                values[cells[0]] = cells[1]

    if values:
        return values

    # 테이블 없으면 문단 텍스트 추출
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise ValueError("DOCX 파일에서 텍스트를 찾지 못했습니다.")

    # "key: value" 또는 "key = value" 패턴 탐지
    for p in paragraphs:
        for sep in [':', '=', '\t']:
            if sep in p:
                parts = p.split(sep, 1)
                if len(parts) == 2 and parts[0].strip():
                    values[parts[0].strip()] = parts[1].strip()
                break

    if values:
        return values

    # 패턴 없으면 전체 텍스트를 하나의 값으로
    return {"내용": '\n'.join(paragraphs)}


def _extract_hwpx_text_as_values(hwpx_path: str) -> dict:
    """HWPX 파일에서 텍스트 추출 → 필드명:값 매핑 시도.

    테이블 구조: 왼쪽 셀=필드명, 오른쪽 셀=값으로 추정.
    """
    try:
        with ZipFile(hwpx_path, 'r') as zf:
            section_xml = zf.read('Contents/section0.xml').decode('utf-8')
    except Exception as e:
        raise ValueError(f"HWPX 읽기 실패: {e}")

    root = etree.fromstring(section_xml.encode('utf-8'))
    ns = {'hp': 'http://www.hancom.co.kr/hwpml/2011/paragraph'}

    # 테이블 셀에서 텍스트 추출
    values = {}
    for tc in root.iter('{http://www.hancom.co.kr/hwpml/2011/paragraph}tc'):
        cell_addr = tc.find('{http://www.hancom.co.kr/hwpml/2011/paragraph}cellAddr')
        if cell_addr is None:
            continue
        col = int(cell_addr.get('colAddr', 0))
        row = int(cell_addr.get('rowAddr', 0))

        texts = []
        for t in tc.iter('{http://www.hancom.co.kr/hwpml/2011/paragraph}t'):
            if t.text:
                texts.append(t.text)
        cell_text = ' '.join(texts).strip()

        values[(row, col)] = cell_text

    # 2열 테이블 → key-value 매핑 (col 0 = 필드명, col 1 = 값)
    result = {}
    rows_seen = set(r for r, c in values.keys())
    for r in sorted(rows_seen):
        key = values.get((r, 0), '').strip()
        val = values.get((r, 1), '').strip()
        if key:
            result[key] = val

    if not result:
        # 2열 구조가 아니면 전체 텍스트를 "내용" 키로
        all_text = []
        for t in root.iter('{http://www.hancom.co.kr/hwpml/2011/paragraph}t'):
            if t.text:
                all_text.append(t.text)
        if all_text:
            result["내용"] = ' '.join(all_text)

    if not result:
        raise ValueError("HWPX 파일에서 텍스트를 추출하지 못했습니다.")
    return result


def _parse_txt_keyvalue(txt_path: str) -> dict:
    """TXT 파일에서 key=value 또는 key: value 쌍 추출"""
    values = {}
    with open(txt_path, 'r', encoding='utf-8-sig') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            for sep in ['=', ':', '\t']:
                if sep in line:
                    parts = line.split(sep, 1)
                    if len(parts) == 2 and parts[0].strip():
                        values[parts[0].strip()] = parts[1].strip()
                    break
    if not values:
        raise ValueError("TXT 파일에서 key=value 또는 key: value 쌍을 찾지 못했습니다.")
    return values


def _xml_escape(text: str) -> str:
    """XML 특수문자 이스케이프"""
    return (text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&apos;"))


def _pack_hwpx(work_dir: Path, output_path: Path) -> None:
    """디렉토리를 HWPX로 패킹 (mimetype first, ZIP_STORED)"""
    mimetype_file = work_dir / 'mimetype'

    all_files = sorted(
        p.relative_to(work_dir).as_posix()
        for p in work_dir.rglob('*')
        if p.is_file()
    )

    with ZipFile(output_path, 'w', ZIP_DEFLATED) as zf:
        if mimetype_file.is_file():
            zf.write(mimetype_file, 'mimetype', compress_type=ZIP_STORED)
        for rel_path in all_files:
            if rel_path == 'mimetype':
                continue
            zf.write(work_dir / rel_path, rel_path, compress_type=ZIP_DEFLATED)
